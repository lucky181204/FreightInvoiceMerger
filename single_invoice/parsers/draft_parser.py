"""Draft parser for single invoice generation.

Extracts fields from BL Draft (.docx) file following defined rules.
Draft contains two tables with shipping document data.

Expected template cells and their format:
  B5  = "Loading Port:{port},{country}"
  B6  = "Destination: {port},{country}"
  B7  = "Shipment of: {qty}x{type}"
  H4  = "Invoice date:{date}"
  H5  = "ETD:{date}"
  H6  = "ATD:{date}"
  H7  = "Invoice #:{business_no}"
  B10 = "PRODUCT: {product_name}"       (Row 8 col 2, first line)
  B11 = "CLASS: {danger_class}"         (Row 8 col 2, parsed)
  B12 = "Shipper: {shipper_name}"       (Row 3 col 2, first line)
  C10 = "PO{po}"
  D10 = "{booking_no}"
  E10 = "{vessel_voyage}"
  F10+ = container_no
  G10+ = container_type
  H10+ = "USD"
  B30 = TOTAL formula (no space after USD)
"""

import logging
import re
from docx import Document

logger = logging.getLogger(__name__)

DESTINATION_COUNTRY_MAP = {
    "SANTOS": "BRAZIL",
    "BUENOS AIRES": "ARGENTINA",
    "BUENAVENTURA": "COLOMBIA",
    "ROTTERDAM": "THE NETHERLANDS",
    "KOPER": "SLOVENIA",
    "VALENCIA": "SPAIN",
    "VILLETA": "PARAGUAY",
}


def get_cell(table, row_idx: int, col_idx: int) -> str:
    """Safely get cell text from a docx table. row_idx/col_idx are 0-based."""
    try:
        return table.cell(row_idx, col_idx).text.strip()
    except Exception:
        return ""


def first_nonempty_line(text: str) -> str:
    """Return the first non-empty line from text, stripped."""
    for line in text.splitlines():
        s = line.strip()
        if s:
            return s
    return ""


def extract_draft_data(draft_path: str) -> dict:
    """
    Extract all fields from BL Draft .docx file.

    Draft Table 1 (booking info) — 9 rows x 4 cols:
      Row 0: OCEAN VESSEL VOYAGE | {vessel}
      Row 1: BOOKING NO. | {booking_no}
      Row 2: SHIPPER: | {shipper_name}         → B12
      Row 5: PORT OF LOADING: | {load_port} | PORT OF DISCHARGE: | {discharge_port}  → B6
      Row 6: MARKS: | DESCRIPTION OF GOODS :
      Row 7: N/M | {goods_desc}                → B10 (first line), B11 (parsed)
      Row 8: SHIPPING TERMS: | {terms}

    Draft Table 2 (container info):
      Row 0: EQU: | {container_info}  (e.g. "1*40HQ")
    """
    doc = Document(draft_path)
    tables = doc.tables
    result = {}

    if len(tables) >= 1:
        t1 = tables[0]

        # E10: Vessel/Voyage (Row 0, col 1)
        result["E10"] = get_cell(t1, 0, 1)

        # D10: Booking No (Row 1, col 1)
        result["D10"] = get_cell(t1, 1, 1)

        # ── B12: Shipper name (Row 2/3, col 1) — first line → PRODUCT field
        # Per spec: B12 = Draft Table 1 Row 3 (0-indexed row 2) Col 2 → Shipper name
        result["shipper_name"] = first_nonempty_line(get_cell(t1, 2, 1))

        # B5 first word (for Loading Port prefix)
        result["B5_first"] = result["shipper_name"].split()[0] if result["shipper_name"] else ""

        # B10: Loading Port (Row 5/PORT OF LOADING, col 1)
        result["B10_loading_port"] = get_cell(t1, 5, 1)

        # ── B6: Discharge Port (Row 5, col 3)
        # Take first line, if comma take before comma, uppercase, then map
        raw_dest = get_cell(t1, 5, 3)
        destination = normalize_destination(raw_dest)
        country = DESTINATION_COUNTRY_MAP.get(destination, "ERROR")
        result["B6_destination"] = destination
        result["B6_country"] = country

        # ── B10: Product name (Row 7/8, col 1) — first line
        desc_text = get_cell(t1, 7, 1)
        result["product_name"] = first_nonempty_line(desc_text)

        # ── B11: Dangerous goods class from description
        result["B11_danger_class"] = extract_cargo_class(desc_text)

    else:
        result["E10"] = ""
        result["D10"] = ""
        result["shipper_name"] = ""
        result["B5_first"] = ""
        result["B10_loading_port"] = ""
        result["B6_destination"] = ""
        result["B6_country"] = ""
        result["product_name"] = ""
        result["B11_danger_class"] = "GENERAL CARGO"

    if len(tables) >= 2:
        t2 = tables[1]
        result["B7_container"] = get_cell(t2, 0, 1)
    else:
        result["B7_container"] = ""

    return result


def normalize_destination(value: str) -> str:
    """
    Extract destination port from raw cell text.
    - First non-empty line
    - If comma, take before comma
    - Strip, uppercase
    """
    first_line = first_nonempty_line(value)
    destination = first_line.split(",", 1)[0].strip().upper()
    return destination


def extract_cargo_class(cell_text: str) -> str:
    """
    Extract cargo class from draft description cell text.

    Priority:
    1. Look for "数字 / 4位数字" pattern (e.g. "9/3077", "6.1 / 2757")
    2. Look for CLASS keyword followed by number (e.g. "CLASS:9", "CLASS: 6.1")
    3. Return "GENERAL CARGO" if neither found.

    Returns the class number only (e.g. "9", "6.1").
    """
    if not cell_text:
        return "GENERAL CARGO"

    lines = [line.strip() for line in cell_text.splitlines() if line.strip()]

    slash_pattern = re.compile(r"(\d+(?:\.\d+)?)\s*/\s*(\d{4})(?!\d)")
    class_pattern = re.compile(
        r"\b(?:CLASS|CL)\b"
        r"\s*[.:：-]?\s*"
        r"(\d+(?:\.\d+)?)",
        re.IGNORECASE
    )

    # Priority 1: exactly one / and /后面的4位数字
    for line in lines:
        if line.count("/") != 1:
            continue
        m = slash_pattern.search(line)
        if m:
            return f"{m.group(1)} HAZ"

    # Priority 2: CLASS or CL keyword
    for line in lines:
        m = class_pattern.search(line)
        if m:
            return f"{m.group(1)} HAZ"

    return "GENERAL CARGO"


def search_po_in_draft(draft_path: str, po: str) -> bool:
    """Search for PO string throughout entire Draft .docx."""
    if not po:
        return False
    doc = Document(draft_path)
    for p in doc.paragraphs:
        if p.text and po in p.text:
            return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if po in cell.text:
                    return True
    return False


def get_container_qty(text: str) -> int:
    """Extract container quantity from text like '1*40HQ' → 1."""
    if not text:
        return 0
    total = 0
    for m in re.finditer(r'(\d+)\s*\*', text):
        total += int(m.group(1))
    return total if total > 0 else 0
